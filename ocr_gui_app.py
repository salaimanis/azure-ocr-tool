import streamlit as st
import io
from PIL import Image
from azure.ai.formrecognizer import DocumentAnalysisClient
from azure.core.credentials import AzureKeyCredential
from docx import Document
from fpdf import FPDF
from openpyxl import Workbook
from datetime import datetime

# Azure credentials
try:
    AZURE_ENDPOINT = st.secrets["azure"]["AZURE_ENDPOINT"]
    AZURE_KEY = st.secrets["azure"]["AZURE_KEY"]
except KeyError:
    st.error("Please configure Azure credentials in `.streamlit/secrets.toml`.")
    st.stop()

# Azure client
client = DocumentAnalysisClient(
    endpoint=AZURE_ENDPOINT,
    credential=AzureKeyCredential(AZURE_KEY)
)

st.set_page_config(page_title="AI OCR Tool", layout="centered")
st.title("📄 AI OCR Tool with Azure Document Intelligence")
st.markdown("Upload an image and extract clean text. Tables supported ✨")

uploaded_file = st.file_uploader("Choose an image", type=["png", "jpg", "jpeg"])
output_format = st.selectbox("Select output format", ["Text", "Word", "PDF", "Excel"])
file_name = st.text_input("Enter output file name (without extension)", value=f"ocr_output_{datetime.now().strftime('%Y%m%d_%H%M%S')}")
model_type = st.radio("Select OCR Model", ["prebuilt-document (Structured)", "prebuilt-read (Handwritten/Text only)"])

if uploaded_file and st.button("Extract Text"):
    image = Image.open(uploaded_file).convert('RGB')
    st.image(image, caption="Uploaded Image", use_container_width=True)

    image_bytes = io.BytesIO()
    image.save(image_bytes, format="JPEG")
    image_bytes.seek(0)

    with st.spinner("Extracting text using Azure OCR..."):
        try:
            model = "prebuilt-document" if "Structured" in model_type else "prebuilt-read"
            poller = client.begin_analyze_document(model, document=image_bytes)
            result = poller.result()

            extracted_text = ""
            tables = []

            for page in result.pages:
                for line in page.lines:
                    extracted_text += line.content + "\n"

            if model == "prebuilt-document":
                for table in result.tables:
                    num_rows = table.row_count
                    num_cols = table.column_count
                    table_content = [["" for _ in range(num_cols)] for _ in range(num_rows)]
                    for cell in table.cells:
                        table_content[cell.row_index][cell.column_index] = cell.content
                    tables.append(table_content)

            if extracted_text:
                st.subheader("Extracted Text")
                st.text_area("Extracted Text Area", extracted_text, height=300, label_visibility="collapsed")

            if tables:
                st.subheader("Detected Table(s)")
                for i, table in enumerate(tables):
                    st.markdown(f"**Table {i+1}:**")
                    st.table(table)

            # Download options
            if output_format == "Text":
                with open(f"{file_name}.txt", "w", encoding="utf-8") as f:
                    f.write(extracted_text)
                with open(f"{file_name}.txt", "rb") as f:
                    st.download_button("Download Text", f, file_name=f"{file_name}.txt")

            elif output_format == "Word":
                doc = Document()
                doc.add_paragraph(extracted_text)
                if tables:
                    for table in tables:
                        word_table = doc.add_table(rows=len(table), cols=len(table[0]))
                        for i, row in enumerate(table):
                            for j, cell in enumerate(row):
                                word_table.cell(i, j).text = cell
                doc.save(f"{file_name}.docx")
                with open(f"{file_name}.docx", "rb") as f:
                    st.download_button("Download Word", f, file_name=f"{file_name}.docx")

            elif output_format == "PDF":
                pdf = FPDF()
                pdf.add_page()
                pdf.set_font("Arial", size=12)
                for line in extracted_text.split("\n"):
                    pdf.multi_cell(0, 10, txt=line)
                if tables:
                    pdf.ln(10)
                    for table in tables:
                        for row in table:
                            row_text = " | ".join(row)
                            pdf.cell(0, 10, txt=row_text, ln=True)
                        pdf.ln(5)
                pdf.output(f"{file_name}.pdf")
                with open(f"{file_name}.pdf", "rb") as f:
                    st.download_button("Download PDF", f, file_name=f"{file_name}.pdf")

            elif output_format == "Excel":
                wb = Workbook()
                ws = wb.active
                ws.title = "Extracted_Text"
                for i, line in enumerate(extracted_text.strip().split("\n")):
                    ws.cell(row=i+1, column=1, value=line)

                for idx, table in enumerate(tables):
                    sheet = wb.create_sheet(title=f"Table_{idx+1}")
                    for r_idx, row in enumerate(table):
                        for c_idx, cell in enumerate(row):
                            sheet.cell(row=r_idx+1, column=c_idx+1, value=cell)

                excel_bytes = io.BytesIO()
                wb.save(excel_bytes)
                excel_bytes.seek(0)
                st.download_button("Download Excel", excel_bytes, file_name=f"{file_name}.xlsx")

            st.success("✅ Done!")

        except Exception as e:
            st.error(f"❌ Error: {str(e)}")
